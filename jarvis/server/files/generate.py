"""Generate real .pdf and .docx files from markdown-ish text."""
from __future__ import annotations

import io
import re


def _blocks(md: str):
    """Parse text into (kind, text) blocks: h1/h2/h3, bullet, number, para."""
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            yield ("blank", "")
            continue
        if line.startswith("### "):
            yield ("h3", line[4:])
        elif line.startswith("## "):
            yield ("h2", line[3:])
        elif line.startswith("# "):
            yield ("h1", line[2:])
        elif re.match(r"^\s*[-*]\s+", line):
            yield ("bullet", re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+[.)]\s+", line):
            yield ("number", re.sub(r"^\s*\d+[.)]\s+", "", line))
        else:
            yield ("para", line)


def _inline(text: str) -> str:
    """Strip markdown emphasis for plain rendering (keep the words)."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def markdown_to_pdf(md: str, title: str = "") -> bytes:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                            leftMargin=2 * cm, rightMargin=2 * cm, title=title or "Document")
    ss = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=ss["Normal"], fontSize=11, leading=16, alignment=TA_LEFT, spaceAfter=6)
    flow = []
    if title:
        flow.append(Paragraph(_inline(title), ParagraphStyle("t", parent=ss["Title"], fontSize=20, spaceAfter=14)))
    pending_list, list_kind = [], None

    def flush_list():
        nonlocal pending_list, list_kind
        if pending_list:
            flow.append(ListFlowable(
                [ListItem(Paragraph(_inline(x), body)) for x in pending_list],
                bulletType="bullet" if list_kind == "bullet" else "1",
            ))
            pending_list, list_kind = [], None

    for kind, text in _blocks(md):
        if kind in ("bullet", "number"):
            if list_kind and list_kind != kind:
                flush_list()
            list_kind = kind
            pending_list.append(text)
            continue
        flush_list()
        if kind == "h1":
            flow.append(Paragraph(_inline(text), ParagraphStyle("h1", parent=ss["Heading1"], fontSize=17, spaceBefore=10, spaceAfter=8)))
        elif kind == "h2":
            flow.append(Paragraph(_inline(text), ParagraphStyle("h2", parent=ss["Heading2"], fontSize=14, spaceBefore=8, spaceAfter=6)))
        elif kind == "h3":
            flow.append(Paragraph(_inline(text), ParagraphStyle("h3", parent=ss["Heading3"], fontSize=12, spaceBefore=6, spaceAfter=4)))
        elif kind == "para":
            flow.append(Paragraph(_inline(text), body))
        elif kind == "blank":
            flow.append(Spacer(1, 4))
    flush_list()
    doc.build(flow)
    return buf.getvalue()


def markdown_to_docx(md: str, title: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    if title:
        doc.add_heading(_inline(title), level=0)
    for kind, text in _blocks(md):
        if kind == "h1":
            doc.add_heading(_inline(text), level=1)
        elif kind == "h2":
            doc.add_heading(_inline(text), level=2)
        elif kind == "h3":
            doc.add_heading(_inline(text), level=3)
        elif kind == "bullet":
            doc.add_paragraph(_inline(text), style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(_inline(text), style="List Number")
        elif kind == "para":
            doc.add_paragraph(_inline(text))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(md: str, fmt: str, title: str = "") -> bytes:
    if fmt == "pdf":
        return markdown_to_pdf(md, title)
    if fmt == "docx":
        return markdown_to_docx(md, title)
    return md.encode("utf-8")  # txt / md
