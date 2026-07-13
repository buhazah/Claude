"""Extract text from uploaded documents so JARVIS can answer questions about
them. Images are handled separately (via the vision model)."""
from __future__ import annotations

import csv
import io
import json
import re

MAX_CHARS = 16000
IMAGE_EXTS = {"png", "jpg", "jpeg", "gif", "webp", "bmp"}
IMAGE_MIMES = {"image/png", "image/jpeg", "image/gif", "image/webp", "image/bmp"}


def is_image(filename: str, content_type: str = "") -> bool:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in IMAGE_EXTS or (content_type or "").lower() in IMAGE_MIMES


def extract_text(data: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    try:
        if ext == "pdf":
            text = _pdf(data)
        elif ext == "docx":
            text = _docx(data)
        elif ext == "json":
            text = json.dumps(json.loads(data.decode("utf-8", "replace")), indent=2)
        elif ext == "csv":
            text = _csv(data)
        elif ext == "rtf":
            text = _rtf(data)
        else:  # txt, md, and anything else readable
            text = data.decode("utf-8", "replace")
    except Exception as exc:  # never crash on a malformed upload
        return f"(Could not read this file: {type(exc).__name__}: {exc})"
    text = text.strip()
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n…(truncated)"
    return text or "(No readable text found in this file.)"


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages)


def _docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _csv(data: bytes) -> str:
    rows = list(csv.reader(io.StringIO(data.decode("utf-8", "replace"))))
    return "\n".join(", ".join(r) for r in rows[:200])


def _rtf(data: bytes) -> str:
    text = data.decode("utf-8", "replace")
    text = re.sub(r"\\[a-z]+-?\d* ?", " ", text)  # strip control words
    text = re.sub(r"[{}]", "", text)
    return re.sub(r"\s+", " ", text).strip()
